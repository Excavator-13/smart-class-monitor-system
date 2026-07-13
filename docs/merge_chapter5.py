import re
import subprocess
import os
import sys
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

sys.stdout.reconfigure(encoding='utf-8')

docx_path = r'D:\team-homework\smart-class-monitor-system\docs\产品需求与设计文档.docx'
md_path = r'D:\team-homework\smart-class-monitor-system\docs\详细设计\第5章_详细设计与实现.md'
temp_dir = r'D:\team-homework\smart-class-monitor-system\docs\mermaid_temp'
npx_cmd = r'D:\software\npx.cmd'

with open(md_path, 'r', encoding='utf-8') as f:
    md_content = f.read()

mermaid_pattern = re.compile(r'```mermaid\n(.*?)\n```', re.DOTALL)
mermaid_blocks = mermaid_pattern.findall(md_content)

print(f"找到 {len(mermaid_blocks)} 个 mermaid 代码块")

mermaid_images = []
for i, code in enumerate(mermaid_blocks):
    mmd_file = os.path.join(temp_dir, f'diagram_{i}.mmd')
    png_file = os.path.join(temp_dir, f'diagram_{i}.png')
    
    with open(mmd_file, 'w', encoding='utf-8') as f:
        f.write(code)
    
    print(f"渲染图表 {i+1}/{len(mermaid_blocks)}...")
    try:
        result = subprocess.run(
            [npx_cmd, '-y', '@mermaid-js/mermaid-cli', '-i', mmd_file, '-o', png_file, '-b', 'white'],
            capture_output=True,
            text=True,
            timeout=60,
            cwd=temp_dir
        )
        if os.path.exists(png_file):
            mermaid_images.append(png_file)
            print(f"  OK 图表 {i+1}")
        else:
            print(f"  FAIL 图表 {i+1}: {result.stderr[:200]}")
            mermaid_images.append(None)
    except Exception as e:
        print(f"  FAIL 图表 {i+1} 异常: {e}")
        mermaid_images.append(None)

success_count = sum(1 for img in mermaid_images if img)
print(f"\n成功渲染 {success_count} 个图表")

def parse_markdown(content):
    lines = content.split('\n')
    paragraphs = []
    in_mermaid = False
    mermaid_index = 0
    
    for line in lines:
        if line.strip() == '```mermaid':
            in_mermaid = True
            continue
        elif in_mermaid and line.strip() == '```':
            in_mermaid = False
            paragraphs.append({'type': 'image', 'index': mermaid_index})
            mermaid_index += 1
            continue
        elif in_mermaid:
            continue
        
        if line.startswith('# '):
            paragraphs.append({'type': 'heading1', 'text': line[2:]})
        elif line.startswith('## '):
            paragraphs.append({'type': 'heading2', 'text': line[3:]})
        elif line.startswith('### '):
            paragraphs.append({'type': 'heading3', 'text': line[4:]})
        elif line.strip() == '':
            paragraphs.append({'type': 'empty', 'text': ''})
        else:
            paragraphs.append({'type': 'text', 'text': line})
    
    return paragraphs

paragraphs = parse_markdown(md_content)

doc = Document(docx_path)

chapter5_start = None
for i, para in enumerate(doc.paragraphs):
    if para.text.strip() == '5. 详细设计与实现':
        chapter5_start = i
        break

if chapter5_start is None:
    print("错误: 找不到第5章")
    exit(1)

print(f"\n第5章位于段落 {chapter5_start}")

body = doc.element.body
target_elements = []
for i, para in enumerate(doc.paragraphs):
    if i >= chapter5_start:
        target_elements.append(para._element)

for elem in target_elements:
    body.remove(elem)

print(f"删除了 {len(target_elements)} 个占位段落")

def add_paragraph(doc, text, style=None):
    para = doc.add_paragraph()
    if style:
        para.style = style
    para.add_run(text)
    return para

def add_image(doc, image_path, width_inches=6.0):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(image_path, width=Inches(width_inches))
    return para

for p in paragraphs:
    if p['type'] == 'heading1':
        add_paragraph(doc, p['text'], 'Heading 1')
    elif p['type'] == 'heading2':
        add_paragraph(doc, p['text'], 'Heading 2')
    elif p['type'] == 'heading3':
        add_paragraph(doc, p['text'], 'Heading 3')
    elif p['type'] == 'image':
        img_path = mermaid_images[p['index']]
        if img_path and os.path.exists(img_path):
            add_image(doc, img_path, 6.0)
        else:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.add_run(f'[图表 {p["index"] + 1} 渲染失败]')
    elif p['type'] == 'text':
        add_paragraph(doc, p['text'])
    elif p['type'] == 'empty':
        doc.add_paragraph()

doc.save(docx_path)

print(f"\n文档已保存到: {docx_path}")
print(f"添加了 {len(paragraphs)} 个段落")
print(f"插入了 {success_count} 个图表")
